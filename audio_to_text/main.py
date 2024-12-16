import subprocess
import os
from dotenv import load_dotenv
from openai import OpenAI
from docx import Document

# Load environment variables
load_dotenv()

# Initialize OpenAI client
client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))

# Audio file to transcription
audio_file = open("audio.mp3", "rb")
transcription = client.audio.transcriptions.create(
    model="whisper-1",
    file=audio_file
)
print(transcription.text)
transcript_text = transcription.text

# Get summary and symptoms analysis
completion = client.chat.completions.create(
    model="gpt-4o-mini",
    messages=[
        {"role": "system", "content": f"""You are an assistant that analyzes transcription and summarizes the patient's symptoms. Use checkboxes (☑ for present, ☐ for absent):
        Transcription:
        {transcript_text}

        Symptoms:
        - Fever
        - Cough
        - Shortness of breath
        - Fatigue
        - Headache

        Mark the checkboxes accordingly.
        """},
        {"role": "user", "content": f"Summarize the following text:\n\n{transcript_text}"}
    ]
)

print("Summary:", completion.choices[0].message.content)
summary = completion.choices[0].message.content

# Save to Word file
doc = Document()
doc.add_heading('Patient Symptoms Summary', level=1)
doc.add_paragraph(summary)

file_name = "Patient_Symptoms_Summary.docx"
doc.save(file_name)
print(f"Document saved as {file_name}")


